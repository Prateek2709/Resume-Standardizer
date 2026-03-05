# firm_resume_docx_only_UPDATED.py
# Correct renderer: render Company template via docxtpl (keeps your firm format)

from docxtpl import DocxTemplate
from jinja2 import Environment

from langsmith import traceable

@traceable(name='render_firm_resume')
def render_firm_resume(*, template_path=None, firm_json=None, out_docx_path=None, output_path=None, **_kwargs):
    if not template_path:
        raise ValueError("template_path is required")
    if firm_json is None:
        raise ValueError("firm_json is required")

    out_path = out_docx_path or output_path
    if not out_path:
        raise ValueError("out_docx_path (or output_path) is required")

    doc = DocxTemplate(template_path)

    jinja_env = Environment(
        trim_blocks=True,
        lstrip_blocks=True,
        autoescape=False
    )

    doc.render(firm_json, jinja_env=jinja_env)
    doc.save(out_path)

def _remove_empty_paragraphs(doc):
    """
    Remove empty paragraphs safely:
    1) Remove empty paragraphs ONLY when they are between two bullet/list paragraphs.
    2) ALSO remove empty paragraphs that occur right after a bullet list and right before
       a new client header (prevents 2-gap between client groups).
    """
    paras = list(doc.paragraphs)
    n = len(paras)

    def _is_bullet(p):
        style_name = (p.style.name or "") if p.style else ""
        has_numbering = p._p.pPr is not None and p._p.pPr.numPr is not None
        return ("List" in style_name) or has_numbering

    def _looks_like_client_header(p):
        txt = (p.text or "").strip()
        if not txt:
            return False
        # don't treat these as "client headers"
        if txt.startswith(("Role:", "Duration:", "Tools", "Tools & Environment")):
            return False
        # don't remove spacing before major section headers
        up = txt.upper()
        if up in {"PROJECT EXPERIENCE", "PROFESSIONAL EXPERIENCE", "EXPERIENCE SUMMARY", "TECHNICAL SKILLS",
                  "EDUCATION", "CERTIFICATIONS", "SKILLS SUMMARY", "PROFESSIONAL SUMMARY"}:
            return False
        return True

    for i, p in enumerate(paras):
        if (p.text or "").strip():
            continue

        prev_p = paras[i - 1] if i > 0 else None
        next_p = paras[i + 1] if i < n - 1 else None

        prev_is_bullet = bool(prev_p and _is_bullet(prev_p))
        next_is_bullet = bool(next_p and _is_bullet(next_p))

        # Case 1: bullet-to-bullet empty line → remove
        if prev_is_bullet and next_is_bullet:
            p._element.getparent().remove(p._element)
            continue

        # Case 2 (NEW): bullet-to-client empty line → remove (reduces 2 gaps to 1)
        if prev_is_bullet and next_p and _looks_like_client_header(next_p):
            p._element.getparent().remove(p._element)
            continue

def _tighten_bullets(doc):
    """
    Remove extra whitespace:
    - Tighten bullet paragraphs
    - Reduce spacing before client headers (between project groups)
    """
    from docx.text.paragraph import Paragraph

    paras = list(doc.paragraphs)

    def _is_bullet(p):
        style_name = (p.style.name or "") if p.style else ""
        has_numbering = p._p.pPr is not None and p._p.pPr.numPr is not None
        return ("List" in style_name) or has_numbering

    for i, p in enumerate(paras):
        style_name = (p.style.name or "") if p.style else ""
        has_numbering = p._p.pPr is not None and p._p.pPr.numPr is not None

        # 1️⃣ Tighten bullets
        if "List" in style_name or has_numbering:
            pf = p.paragraph_format
            pf.space_before = 0
            pf.space_after = 0
            pf.line_spacing = 1.0

        # 2️⃣ Reduce spacing before a new client block
        if i > 0:
            prev = paras[i - 1]

            # If previous paragraph was a bullet
            prev_is_bullet = _is_bullet(prev)

            # And current paragraph looks like a client name
            # (client name paragraphs are plain text, not starting with Role:/Duration:)
            text = (p.text or "").strip()
            if prev_is_bullet and text and not text.startswith(("Role:", "Duration:", "Tools")):
                pf = p.paragraph_format
                pf.space_before = 0   # keep ONE line gap (not zero)
                pf.space_after = 0

def _tighten_section_headers(doc, *, keep_space_before_headers: bool = False):
    """
    Remove the 1-line gap between SECTION HEADER (e.g., 'EXPERIENCE SUMMARY')
    and its first content paragraph, by:
    - removing empty paragraph immediately after header
    - setting header space_after = 0
    - setting next paragraph space_before = 0
    """
    paras = list(doc.paragraphs)

    SECTION_HEADERS = {
        "EXPERIENCE SUMMARY",
        "TECHNICAL SKILLS",
        "EDUCATION",
        "CERTIFICATIONS",
        "PROJECT EXPERIENCE",
        "PROFESSIONAL EXPERIENCE",
        "SKILLS SUMMARY",
        "PROFESSIONAL SUMMARY",
    }

    i = 0
    while i < len(paras) - 1:
        p = paras[i]
        txt = (p.text or "").strip().upper()

        if txt in SECTION_HEADERS:
            # 1) kill spacing on the header itself
            pf = p.paragraph_format
            if not keep_space_before_headers:
                pf.space_before = 0
            # else: keep whatever the template has set (Non-Company uses this)

            pf.space_after = 0
            pf.line_spacing = 1.0

            # 2) remove empty paragraphs immediately after the header
            j = i + 1
            while j < len(paras) and not (paras[j].text or "").strip():
                empty_p = paras[j]
                empty_p._element.getparent().remove(empty_p._element)
                # refresh list after mutation
                paras = list(doc.paragraphs)
                j = i + 1

            # 3) tighten the first content paragraph after the header
            if i + 1 < len(paras):
                next_p = paras[i + 1]
                npf = next_p.paragraph_format
                npf.space_before = 0
                # don't force space_after here; let content style control it

            # refresh after edits
            paras = list(doc.paragraphs)

        i += 1

def _normalize_gap_before_client_headers(doc, *, keep_blank_lines: int = 1):
    """
    Ensure EXACTLY `keep_blank_lines` empty paragraphs before each client header,
    and remove visual extra gap by forcing:
      - previous content paragraph: space_after = 0
      - client header paragraph: space_before = 0
    This targets the '2-line gap between previous client summary and next client name'.
    """
    paras = list(doc.paragraphs)

    SECTION_HEADERS = {
        "EXPERIENCE SUMMARY", "TECHNICAL SKILLS", "EDUCATION", "CERTIFICATIONS",
        "PROJECT EXPERIENCE", "PROFESSIONAL EXPERIENCE", "SKILLS SUMMARY", "PROFESSIONAL SUMMARY",
    }

    def _is_section_header(txt: str) -> bool:
        return (txt or "").strip().upper() in SECTION_HEADERS

    def _looks_like_client_header(p):
        txt = (p.text or "").strip()
        if not txt:
            return False
        if txt.startswith(("Role:", "Duration:", "Tools", "Tools & Environment")):
            return False
        if _is_section_header(txt):
            return False
        return True

    def _remove_para(p):
        p._element.getparent().remove(p._element)

    i = 0
    while i < len(paras):
        p = paras[i]
        if _looks_like_client_header(p):
            # Count empty paragraphs immediately before this client header
            empty_idxs = []
            j = i - 1
            while j >= 0 and not (paras[j].text or "").strip():
                empty_idxs.append(j)
                j -= 1

            # Keep only N empties, remove the rest
            # Example: if we want exactly 1 blank line and we have 2+ empties, remove extras.
            if len(empty_idxs) > keep_blank_lines:
                # remove the extras closest to the client header? (doesn't matter)
                to_remove = empty_idxs[keep_blank_lines:]
                for ridx in to_remove:
                    _remove_para(paras[ridx])

                # refresh list after mutation
                paras = list(doc.paragraphs)
                # recompute current i (client header moved up)
                i = min(i, len(paras) - 1)
                p = paras[i]

            # Remove *visual* extra gap caused by spacing:
            # previous non-empty paragraph should not add extra space_after
            # client header should not add space_before
            # find previous non-empty paragraph again
            k = i - 1
            while k >= 0 and not (paras[k].text or "").strip():
                k -= 1
            if k >= 0:
                prev_content = paras[k]
                prev_content.paragraph_format.space_after = 0
                # ❌ DO NOT force space_after=0 here (breaks section gaps in Non-Company template)

            p.paragraph_format.space_before = 0
            p.paragraph_format.space_after = 0
            p.paragraph_format.line_spacing = 1.0

            # refresh list (safe)
            paras = list(doc.paragraphs)

        i += 1

@traceable(name='postprocess_docx_spacing')
def postprocess_docx_spacing(docx_path: str, *, keep_space_before_headers: bool = False):
    from docx import Document
    doc = Document(docx_path)
    _remove_empty_paragraphs(doc)
    _tighten_bullets(doc)
    _tighten_section_headers(doc, keep_space_before_headers=keep_space_before_headers)
    _normalize_gap_before_client_headers(doc, keep_blank_lines=1)
    doc.save(docx_path)