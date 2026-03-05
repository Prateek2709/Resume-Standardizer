# import libraries
from dotenv import load_dotenv
import os
from pathlib import Path
import uuid
import io
import json
import hashlib
import re
import time
from datetime import datetime
from dateutil import parser as date_parser

from openai import AzureOpenAI
import streamlit as st
import streamlit.components.v1 as components
import pandas as pd
import pdfplumber
from openpyxl import load_workbook
from openpyxl.styles import Font, Alignment

from langsmith import traceable

from firm_resume_docx_only import render_firm_resume, postprocess_docx_spacing
from database import insert_resume_upload

# load env variables
load_dotenv()

AZURE_OPENAI_API_KEY = os.getenv("AZURE_OPENAI_API_KEY")
AZURE_OPENAI_ENDPOINT = os.getenv("AZURE_OPENAI_ENDPOINT")
AZURE_OPENAI_API_DEPLOYMENT_NAME = os.getenv("AZURE_OPENAI_API_DEPLOYMENT_NAME")
AZURE_OPENAI_API_VERSION = os.getenv("AZURE_OPENAI_API_VERSION")
os.environ["LANGCHAIN_PROJECT"] = "Resume-Standardizer"

client = AzureOpenAI(
    azure_endpoint=AZURE_OPENAI_ENDPOINT,
    api_key=AZURE_OPENAI_API_KEY,
    api_version=AZURE_OPENAI_API_VERSION,
    timeout=120.0,
    max_retries=2,
)

OUTPUT_ROOT = Path("output")  # root output folder

# Streamlit reruns the script often; keep run folders stable per uploaded file in a session
if "run_map" not in st.session_state:
    st.session_state.run_map = {}  # file_id -> run_dir (Path)
if "processed" not in st.session_state:
    st.session_state.processed = set()  # file_id already fully processed
if "batch_llm_usage" not in st.session_state:
    st.session_state.batch_llm_usage = []  # list of usage dicts for ALL resumes in this run

# Used to force-reset the file_uploader widget (clears all selected files)
if "uploader_key" not in st.session_state:
    st.session_state.uploader_key = str(uuid.uuid4())

TEMPLATE_DOCX = "templates/Company_Template.docx"
NON_TEMPLATE_DOCX = "templates/Non_Company_Template.docx"

st.set_page_config(page_title="AI Resume Parser", layout="wide")

st.markdown("""
<style>
/* --- Uploader: reduce padding/height so it doesn't feel massive --- */
div[data-testid="stFileUploaderDropzone"] {
    padding: 0.75rem 1rem !important;
    min-height: 92px !important;
    border-radius: 14px !important;
}

/* Make the uploader's internal icon a bit smaller */
div[data-testid="stFileUploaderDropzone"] svg {
    width: 30px !important;
    height: 30px !important;
}

/* --- Clear button: make it look like a square tile --- */
button[kind="secondary"][data-testid="baseButton-secondary"] {
    border-radius: 22px !important;
}

/* --- Clear tile wrapper: force squarish look + centered --- */
.clear-tile div[data-testid="stButton"] button {
    height: 110px !important;          /* match uploader min-height */
    min-height: 110px !important;
    width: 300px !important;           /* force square */
    min-width: 300px !important;
    max-width: 300px !important;

    font-weight: 800 !important;
    font-size: 22px !important;
    line-height: 1.05 !important;
    white-space: pre-line !important; /* keep the line break */

    border: 3px solid #111 !important;
    border-radius: 28px !important;   /* rounded square like your image */
    padding: 0 !important;

    display: flex !important;
    align-items: center !important;
    justify-content: center !important;
    text-align: center !important;
}
</style>
""", unsafe_allow_html=True)


# ------------------------- Helpers -------------------------

def ensure_dir(p: Path) -> Path:
    p.mkdir(parents=True, exist_ok=True)
    return p


def safe_name(name: str, max_len: int = 60) -> str:
    # Keep it filesystem-safe and readable
    stem = Path(name).stem
    stem = re.sub(r"[^A-Za-z0-9_-]+", "_", stem).strip("_")
    return (stem[:max_len] or "resume")


def create_run_dir(uploaded_filename: str) -> Path:
    """
    Creates a unique subfolder under output/ for a single uploaded resume.
    Example: output/20251229_153012_BennyRajaA_9f3a1c2d/
    """
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    job_id = str(uuid.uuid4())[:8]
    sname = safe_name(uploaded_filename)
    run_dir = OUTPUT_ROOT / f"{sname}_{timestamp}_{job_id}"
    return ensure_dir(run_dir)


def _split_tools_env(s: str) -> list[str]:
    """
    Split 'Tools & Environment' strings into individual skills/tools.
    Handles commas, ampersands, pipes.
    """
    if not s:
        return []
    s = str(s)
    # normalize separators
    s = s.replace("|", ",").replace("&", ",")
    parts = [p.strip() for p in s.split(",")]
    return [p for p in parts if p]


def _months_between(start_dt: datetime, end_dt: datetime) -> int:
    """Whole months between two datetimes (at least 1 if same month)."""
    if end_dt < start_dt:
        return 0
    months = (end_dt.year - start_dt.year) * 12 + (end_dt.month - start_dt.month)
    return max(1, months + 1)  # +1 to count the start month


def _parse_duration_to_months(duration: str) -> int:
    """
    Parse strings like:
      'Jan 2024 – Till Date'
      'Jan 2021 – Dec 2023'
      'Sep 2019 – Dec 2021'
    Returns months (int). If can't parse, returns 0.
    """
    if not duration:
        return 0

    s = str(duration)
    s = s.replace("–", "-").replace("—", "-")
    parts = [p.strip() for p in s.split("-") if p.strip()]
    if len(parts) < 2:
        return 0

    start_s, end_s = parts[0], parts[1]
    # normalize end date
    end_lower = end_s.lower()
    if "till date" in end_lower or "present" in end_lower or "current" in end_lower:
        end_dt = datetime.now()
    else:
        try:
            end_dt = date_parser.parse(end_s, default=datetime(1900, 1, 1))
        except Exception:
            return 0

    try:
        start_dt = date_parser.parse(start_s, default=datetime(1900, 1, 1))
    except Exception:
        return 0

    return _months_between(start_dt, end_dt)


def _months_to_label(months: int) -> str:
    if months <= 0:
        return ""
    if months < 12:
        return "Less than 1 year"
    years = months // 12
    return "1 year" if years == 1 else f"{years} years"


def _normalize_availability(v) -> str:
    s = "" if v is None else str(v).strip().lower()
    return "Available" if s in {"yes", "y", "available", "true"} else NOT_AVAILABLE


def build_dynamic_skill_matrix(firm_json: dict) -> pd.DataFrame:
    """
    Build a dynamic skill matrix from firm_json projects:
    - skill appears in tools_environment => assign role duration
    - if repeated across roles => keep max duration (months)
    """
    projects = firm_json.get("projects", []) or []
    best_months_by_skill = {}
    display_name_by_key = {}

    for p in projects:
        if not isinstance(p, dict):
            continue
        tools_env = p.get("tools_environment", "") or p.get("tools", "") or ""
        duration = p.get("duration", "") or ""
        months = _parse_duration_to_months(duration)

        for skill in _split_tools_env(tools_env):
            key = skill.strip().lower()
            if not key:
                continue
            # keep a stable display name (first seen)
            if key not in display_name_by_key:
                display_name_by_key[key] = skill.strip()

            prev = best_months_by_skill.get(key, 0)
            best_months_by_skill[key] = max(prev, months)

    rows = []
    for key, months in best_months_by_skill.items():
        label = _months_to_label(months)
        if label:  # skip anything we couldn't compute
            rows.append({"Skill": display_name_by_key.get(key, key), "Experience": label})

    df = pd.DataFrame(rows)
    if not df.empty:
        def _label_to_months(label: str) -> int:
            if not label:
                return 0
            if label.strip().lower() == "less than 1 year":
                return 11
            m = re.match(r"^\s*(\d+)\s+year", label.strip().lower())
            if m:
                return int(m.group(1)) * 12
            return 0

        df["_months"] = df["Experience"].apply(_label_to_months)
        df = df.sort_values(by=["_months", "Skill"], ascending=[False, True]).drop(columns=["_months"]).reset_index(drop=True)
    return df


@traceable(name='load_pdf_bytes')
def load_pdf_bytes(pdf_bytes: bytes) -> str:
    """Extract text from a PDF provided as bytes (no file saved to disk)."""
    text = ""
    with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            text += page_text + "\n"
    return text


@traceable(name='load_docx_bytes')
def load_docx_bytes(docx_bytes: bytes) -> str:
    """Extract text from a DOCX resume provided as bytes (includes headers, footers, and tables)."""
    from docx import Document
    doc = Document(io.BytesIO(docx_bytes))

    chunks: list[str] = []

    # Body paragraphs
    for para in doc.paragraphs:
        t = (para.text or "").strip()
        if t:
            chunks.append(t)

    # Tables (common for contact info, skills grids, etc.)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                ct = (cell.text or "").strip()
                if ct:
                    chunks.append(ct)

    # Headers / Footers (contact info is often here)
    for section in doc.sections:
        for hp in section.header.paragraphs:
            ht = (hp.text or "").strip()
            if ht:
                chunks.append(ht)
        for fp in section.footer.paragraphs:
            ft = (fp.text or "").strip()
            if ft:
                chunks.append(ft)

    # De-dup while preserving order (headers/tables can repeat)
    seen = set()
    out = []
    for c in chunks:
        if c not in seen:
            seen.add(c)
            out.append(c)

    return "\n".join(out)


def _reset_batch_llm_usage():
    st.session_state.batch_llm_usage = []


def _print_batch_llm_totals():
    rows = st.session_state.get("batch_llm_usage", []) or []

    total_inp = sum(int(r.get("input_tokens", 0) or 0) for r in rows)
    total_out = sum(int(r.get("output_tokens", 0) or 0) for r in rows)
    total_all = sum(int(r.get("total_tokens", 0) or 0) for r in rows)
    total_time = round(sum(float(r.get("response_time_sec", 0) or 0) for r in rows), 2)

    print("\n===================== LLM USAGE (COMBINED) =====================")
    print(f"TOTAL INPUT tokens : {total_inp}")
    print(f"TOTAL OUTPUT tokens: {total_out}")
    print(f"TOTAL TOKENS       : {total_all}")
    print(f"TOTAL TIME (sec)   : {total_time}")
    print(f"TOTAL LLM CALLS    : {len(rows)}")
    print("=====================================================================\n")

# ------------------------- UI helpers: Copy table -------------------------

def _df_to_styled_html(df: pd.DataFrame) -> str:
    """Return an HTML table with borders + header styling (easy to paste into Word/Email)."""
    if df is None:
        df = pd.DataFrame()
    html = df.to_html(index=False, escape=False)
    # add simple inline CSS so borders come along when pasted
    css = """
    <style>
    table { border-collapse: collapse; font-family: Arial, sans-serif; font-size: 12px; }
    th, td { border: 1px solid #888; padding: 6px 8px; vertical-align: top; }
    th { font-weight: 700; background: #f2f2f2; }
    </style>
    """
    return css + html

def render_copy_button_for_df(df: pd.DataFrame, *, button_label: str, unique_key: str):
    """Renders a one-click 'Copy table' button that copies HTML + plain text to clipboard."""
    # Unique DOM ids per render
    dom_id = f"copytbl_{unique_key}"
    html_table = _df_to_styled_html(df)
    plain = df.to_csv(index=False)

    # Use Clipboard API with HTML (best for Word/Outlook) + fallback to plain text
    comp = f"""
    <div style="margin: 6px 0 14px 0;">
    <button id="{dom_id}_btn" style="padding:6px 10px; border:1px solid #ccc; border-radius:6px; cursor:pointer;">
        {button_label}
    </button>
    <span id="{dom_id}_msg" style="margin-left:10px; font-size:12px;"></span>
    <div id="{dom_id}_html" style="display:none;">{html_table}</div>
    </div>

    <script>
    (async function() {{
    const btn = document.getElementById("{dom_id}_btn");
    const msg = document.getElementById("{dom_id}_msg");
    const html = document.getElementById("{dom_id}_html").innerHTML;
    const text = {plain!r};

    btn.addEventListener("click", async () => {{
        msg.textContent = "";
        try {{
        if (navigator.clipboard && window.ClipboardItem) {{
            const blobHtml = new Blob([html], {{ type: "text/html" }});
            const blobText = new Blob([text], {{ type: "text/plain" }});
            const item = new ClipboardItem({{ "text/html": blobHtml, "text/plain": blobText }});
            await navigator.clipboard.write([item]);
        }} else if (navigator.clipboard) {{
            // fallback: plain text only
            await navigator.clipboard.writeText(text);
        }} else {{
            throw new Error("Clipboard API not available");
        }}
        msg.textContent = "Copied ✅";
        }} catch (e) {{
        console.error(e);
        msg.textContent = "Copy blocked by browser. Try running on https / localhost.";
        }}
    }});
    }})();
    </script>
    """
    components.html(comp, height=70)


# ------------------------- Education heuristics -------------------------

_DEGREE_MARKERS_RE = re.compile(
    r"""(?ix)
    \b(
        bachelor|b\.?\s?tech|b\.?\s?e\b|b\.?\s?sc\b|bca\b|bba\b|
        master|m\.?\s?tech|m\.?\s?e\b|m\.?\s?sc\b|mca\b|mba\b|
        phd\b|doctorate\b|d\.?\s?phil\b
    )\b
    """
)

_YEAR_RE = re.compile(r"\b(19\d{2}|20\d{2})\b")

def _heuristic_extract_education_entries(resume_text: str, max_items: int = 3) -> list[dict]:
    """
    Best-effort extraction for education when multiple degrees are present in one line.
    Does NOT guess missing parts. Returns list of dicts: degree/institution/year_range.
    """
    if not resume_text:
        return []

    t = str(resume_text)
    # Try to capture only the education "line" first (stops bleeding into next sections)
    m = re.search(r"(?im)^\s*education\s*[:\-]?\s*(.+)$", t)
    if m:
        chunk = m.group(1).strip()
    else:
        # Fallback: grab a small window but stop at common next-section headers
        m2 = re.search(r"(?is)\beducation\b\s*[:\-]?\s*(.{0,800})", t)
        chunk = (m2.group(1) if m2 else t[:1200]).strip()
        chunk = re.split(r"(?i)\b(experience|work experience|skills|projects|summary|certifications)\b", chunk)[0].strip()

    # Normalize spacing (but preserve separators lightly)
    chunk = re.sub(r"\s+", " ", chunk).strip()

    # Split into potential degree segments by looking ahead for degree markers
    parts = []
    idxs = [m.start() for m in _DEGREE_MARKERS_RE.finditer(chunk)]
    if idxs:
        idxs.append(len(chunk))
        for a, b in zip(idxs, idxs[1:]):
            seg = chunk[a:b].strip(" ;,|-\n\t")
            # ✅ Split a segment if another degree marker appears after a YEAR
            # Example: "... - 2015 Bachelor’s in ..."
            subparts = re.split(r"(?i)(?<=\b(19\d{2}|20\d{2})\b)\s+(?=(bachelor|master|phd|doctorate|b\.|m\.))", seg)
            # re.split with capture groups returns extra items; rebuild cleanly
            rebuilt = []
            buf = ""
            for token in subparts:
                if not token:
                    continue
                # If token looks like a degree marker starting, flush buf
                if _DEGREE_MARKERS_RE.search(token) and buf.strip():
                    rebuilt.append(buf.strip())
                    buf = token
                else:
                    buf = (buf + " " + token).strip() if buf else token
            if buf.strip():
                rebuilt.append(buf.strip())

            for rp in rebuilt:
                rp = (rp or "").strip()
                if rp:
                    parts.append(rp)

            # ❌ IMPORTANT: do NOT append `seg` again (it causes duplicates)
            # if seg:
            #     parts.append(seg)

    else:
        # Fallback: split by common separators if no markers found
        parts = re.split(r"\s{2,}|\s*\|\s*|\s*;\s*", chunk)

    entries = []
    for seg in parts:
        if not seg or len(entries) >= max_items:
            continue

        # Extract year (single year, or year range if two years appear)
        years = _YEAR_RE.findall(seg)
        year_range = ""
        if len(years) >= 2:
            year_range = f"{years[0]}-{years[1]}"
        elif len(years) == 1:
            year_range = years[0]

        # Remove years for cleaner parsing
        seg_wo_year = _YEAR_RE.sub("", seg)
        seg_wo_year = re.sub(r"\s*[-–—]\s*$", "", seg_wo_year).strip()

        # Split on commas: Degree, Institution, Location...
        bits = [b.strip() for b in seg_wo_year.split(",") if b.strip()]

        degree = bits[0] if bits else ""
        institution = ""
        if len(bits) >= 2:
            institution = bits[1]
        # If still empty, attempt: last 'in/at/of' phrase
        if not institution:
            m2 = re.search(r"(?i)\b(?:at|from)\b\s+([^,\-]{3,80})", seg_wo_year)
            if m2:
                institution = m2.group(1).strip()

        # Keep only if something meaningful exists
        if degree:
            entries.append({
                "degree": degree,
                "institution": institution,
                "year_range": year_range
            })

    return entries

# ---------------------------------------------------------------------
# ------------------------- LLM call safeguards -------------------------

def _truncate_for_llm(text: str, max_chars: int = 22000) -> str:
    """Keep prompt size bounded to reduce timeouts / token overflows."""
    if not text:
        return ""
    # normalize whitespace a bit to save tokens
    t = re.sub(r"[ \t]+", " ", str(text))
    t = re.sub(r"\n{3,}", "\n\n", t).strip()
    if len(t) <= max_chars:
        return t
    # keep head + middle + tail so we don't drop anything in the middle
    head_len = int(max_chars * 0.4)
    mid_len  = int(max_chars * 0.2)
    tail_len = max_chars - head_len - mid_len

    head = t[:head_len]
    mid_start = max(0, (len(t) // 2) - (mid_len // 2))
    mid = t[mid_start: mid_start + mid_len]
    tail = t[-tail_len:] if tail_len > 0 else ""

    return (
        head
        + "\n\n...[TRUNCATED_HEAD]...\n\n"
        + mid
        + "\n\n...[TRUNCATED_MIDDLE]...\n\n"
        + tail
    )


# --- Canonical "full extraction" schema (Stage A) ---
FULL_SCHEMA = {
    "document_meta": {
        "detected_name": "",
        "detected_location": "",
        "detected_title": ""
    },
    "sections": [
        {
            "title": "",
            "type": "",  # summary | skills | experience | education | certifications | projects | other
            "order": 0,
            "items": [
                {
                    "kind": "",   # paragraph | bullet | table_row | project | role_block | certification | education_entry | skill_group | other
                    "text": "",
                    "fields": {},
                    "provenance": {
                        "source_section_title": "",
                        "source_section_type": "",
                        "section_order": 0
                    }
                }
            ]
        }
    ]
}


def _chunk_text_for_llm(text: str, chunk_chars: int = 14000, overlap: int = 1200) -> list[str]:
    """
    Split text into overlapping chunks so we never drop the middle.
    chunk_chars chosen to stay comfortably under context limits once prompt is added.
    """
    if not text:
        return []
    t = re.sub(r"[ \t]+", " ", str(text))
    t = re.sub(r"\n{3,}", "\n\n", t).strip()

    chunks = []
    i = 0
    n = len(t)
    while i < n:
        j = min(i + chunk_chars, n)
        chunks.append(t[i:j])
        if j >= n:
            break
        i = max(0, j - overlap)
    return chunks


def _merge_extracted_full(parts: list[dict]) -> dict:
    """
    Merge multiple extracted_full JSONs into one.
    Keeps section order by concatenating; also normalizes order indices.
    """
    out = {"document_meta": {"detected_name": "", "detected_location": "", "detected_title": ""}, "sections": []}

    # document_meta: keep first non-empty seen
    for p in parts:
        dm = (p or {}).get("document_meta") or {}
        for k in ["detected_name", "detected_location", "detected_title"]:
            if not out["document_meta"].get(k) and dm.get(k):
                out["document_meta"][k] = dm.get(k)

    # sections: just append in the order chunks were processed
    sec_order = 0
    for p in parts:
        secs = (p or {}).get("sections") or []
        for s in secs:
            if not isinstance(s, dict):
                continue
            s2 = dict(s)
            s2["order"] = sec_order
            # also fix provenance.section_order inside items (optional but nice)
            items = s2.get("items") or []
            if isinstance(items, list):
                for it in items:
                    if isinstance(it, dict):
                        prov = it.get("provenance") or {}
                        if isinstance(prov, dict):
                            prov["section_order"] = sec_order
                            it["provenance"] = prov
            out["sections"].append(s2)
            sec_order += 1

    return out


@traceable(name='chat_completion_with_retry')
def _chat_completion_with_retry(
    messages,
    *,
    temperature: float = 0,
    max_attempts: int = 4,
    run_dir: Path | None = None,
    call_name: str = "UNNAMED"
):
    last_err = None

    for attempt in range(1, max_attempts + 1):
        try:
            start_time = time.time()

            response = client.chat.completions.create(
                model=AZURE_OPENAI_API_DEPLOYMENT_NAME,
                messages=messages,
                temperature=temperature
            )

            duration = round(time.time() - start_time, 2)

            # Token usage (INPUT = prompt_tokens, OUTPUT = completion_tokens)
            usage_data = {
                "call_name": call_name,
                "input_tokens": 0,
                "output_tokens": 0,
                "total_tokens": 0,
                "response_time_sec": duration,
                "timestamp": datetime.now().isoformat(),
            }

            if getattr(response, "usage", None):
                usage_data["input_tokens"] = int(getattr(response.usage, "prompt_tokens", 0) or 0)
                usage_data["output_tokens"] = int(getattr(response.usage, "completion_tokens", 0) or 0)
                usage_data["total_tokens"] = int(getattr(response.usage, "total_tokens", 0) or (usage_data["input_tokens"] + usage_data["output_tokens"]))

            # ---- PRINT (per call) ----
            print(f"\n========== LLM CALL (Stage): {call_name} ==========")
            print(f"INPUT tokens : {usage_data['input_tokens']}")
            print(f"OUTPUT tokens: {usage_data['output_tokens']}")
            print(f"TOTAL tokens : {usage_data['total_tokens']}")
            print(f"Time (sec)   : {usage_data['response_time_sec']}")
            print("===============================================\n")

            # ---- Save per-resume usage ----
            if run_dir:
                ensure_dir(run_dir)
                usage_path = run_dir / "llm_usage.json"
                existing = []
                if usage_path.exists():
                    try:
                        existing = json.loads(usage_path.read_text(encoding="utf-8"))
                    except Exception:
                        existing = []
                existing.append(usage_data)
                usage_path.write_text(json.dumps(existing, indent=2), encoding="utf-8")

                # Save raw output for debugging
                raw_output_path = run_dir / f"{call_name}_raw_output.json"
                raw_output_path.write_text(response.choices[0].message.content, encoding="utf-8")

            # ---- Save batch-wide usage (ALL resumes combined) ----
            try:
                st.session_state.batch_llm_usage.append(usage_data)
            except Exception:
                pass

            return response

        except Exception as e:
            last_err = e
            time.sleep(min(2 ** (attempt - 1), 8))

    raise last_err

# ---------------------------------------------------------------------


# ------------------------- Safety: DOCX/XML sanitization -------------------------

_XML_CTRL_RE = re.compile(r"[\x00-\x08\x0B\x0C\x0E-\x1F]")

def _sanitize_for_docx(value):
    """Recursively sanitize strings so DOCX XML rendering won't break (e.g., stray '&' or control chars)."""
    if value is None:
        return value
    if isinstance(value, str):
        s = _XML_CTRL_RE.sub("", value)
        s = re.sub(r"&(?![A-Za-z]+;|#\d+;|#x[0-9A-Fa-f]+;)", "and", s)
        s = s.replace("<", "").replace(">", "")
        return s
    if isinstance(value, list):
        return [_sanitize_for_docx(v) for v in value]
    if isinstance(value, dict):
        return {k: _sanitize_for_docx(v) for k, v in value.items()}
    return value

# ------------------------- LLM Call #2: UI/Table enrichment -------------------------

@traceable(name='enrich_ui_fields_via_llm')
def enrich_ui_fields_via_llm(parsed_data: dict, resume_text: str, run_dir: Path, template_df: pd.DataFrame | None = None) -> dict:
    """Second LLM call focused on UI table fields (education + screening fields)."""
    resume_text = _truncate_for_llm(resume_text)
    # Carry forward a robust certifications-present flag for availability-only display.
    parsed_data["_cert_present"] = _cert_present_from_text(resume_text)

    particulars = []
    if template_df is not None and "Particular" in template_df.columns:
        for p in template_df["Particular"].astype(str).tolist():
            p = p.strip()
            if p:
                particulars.append(p)
    particulars = particulars[:120]

    prompt = f"""
    You are extracting ONLY the following from the RESUME_TEXT. Return ONLY valid JSON.

    STRICT RULES:
    - Do NOT guess or infer.
    - If a field is not explicitly present, output exactly "Not available".
    - education_entries: extract up to 3 education items AS WRITTEN (degree + institution + year_range if present).

    OUTPUT JSON SCHEMA (ONLY these keys):
    {{
    "contact_number": "Not available",
    "email": "Not available",
    "linkedin": "Not available",
    "visa_status": "Not available",
    "relocation": "Not available",
    "availability": "Not available",
    "interview_availability": "Not available",
    "currently_on_project": "Not available",
    "any_certification": "Not available",
    "education_entries": [
        {{"degree":"","institution":"","year_range":""}}
    ],
    "education": "Not available",
    "skill_matrix": {{}}
    }}

    skill_matrix:
    - ONLY fill keys that are explicitly present in resume and match any of these labels (exact key names):
    {json.dumps(particulars, ensure_ascii=False)}

    RESUME_TEXT:
    ---
    {resume_text}
    ---
    """

    response = _chat_completion_with_retry(
        messages=[
            {"role": "system", "content": "You are an HR expert analysing resumes."},
            {"role": "user", "content": prompt}
        ],
        temperature=0,
        run_dir=run_dir,
        call_name="enrich_ui_fields_via_llm() — screening fields + education_entries + skill_matrix"
    )

    content = response.choices[0].message.content.strip()
    content = re.sub(r"```(?:json)?", "", content).strip()
    content = content.strip().lstrip("\ufeff").strip()

    try:
        enrich = json.loads(content)
    except json.JSONDecodeError:
        enrich = {}

    # Normalize education_entries
    ee = enrich.get("education_entries")
    if not isinstance(ee, list) or not ee:
        enrich["education_entries"] = [{"degree": "", "institution": "", "year_range": ""}]
    else:
        cleaned = []
        for item in ee[:3]:
            if not isinstance(item, dict):
                continue
            cleaned.append({
                "degree": str(item.get("degree", "") or "").strip(),
                "institution": str(item.get("institution", "") or "").strip(),
                "year_range": str(item.get("year_range", "") or "").strip(),
            })
        enrich["education_entries"] = cleaned or [{"degree": "", "institution": "", "year_range": ""}]

    # Enforce required defaults
    for k in [
        "visa_status","relocation","availability","interview_availability",
        "currently_on_project"
    ]:
        enrich[k] = _required_default(enrich.get(k))

    # any_certification must be availability-only
    enrich["any_certification"] = _normalize_availability(enrich.get("any_certification"))

    for k in ["contact_number","email","linkedin","education"]:
        if k in enrich:
            enrich[k] = _required_default(enrich.get(k))

    if not isinstance(enrich.get("skill_matrix"), dict):
        enrich["skill_matrix"] = {}

    return enrich


@traceable(name='parse_resume')
def parse_resume(text: str, run_dir: Path) -> dict:
    text = _truncate_for_llm(text)
    prompt = f"""
    You are an HR expert with extensive experience analyzing and extracting key information from resumes. Your task is to extract and organize the following details from the provided resume text:
        1. **Name**: The full name of the candidate.
        2. **Role**: The current or most recent role/position of the candidate.
        3. **Location**: The candidate's location or the location mentioned in the work experience or contact information.
        4. **Years of Experience**: If explicitly mentioned, use it. Otherwise, calculate it based on the work experience section by identifying the earliest and latest job roles.
        5. **Certifications**: List any certifications or professional qualifications mentioned.
        6. **Skills**: Extract and categorize skills from **anywhere in the resume** (not just the Skills section). Categorize them into relevant buckets:
            - **Programming Languages**: (e.g., Python, Java, C++, JavaScript)
            - **Databases**: (e.g., MySQL, PostgreSQL, MongoDB, Oracle)
            - **Cloud Platforms**: (e.g., AWS, Azure, GCP)
            - **Project Management Tools**: (e.g., Jira, Trello, Asana)
            - **Reporting & Business Intelligence Tools**: (e.g., Tableau, Power BI, Google Data Studio)
            - **Technologies & Frameworks**: (e.g., React, Django, Spark, Salesforce)
            - **DevOps & CI/CD Tools**: (e.g., Docker, Kubernetes, Jenkins, GitHub Actions)
            - **Productivity Tools**: (e.g., Microsoft Word, PowerPoint, Excel, Notion, Slack)
            - **Soft Skills**: (e.g., Leadership, Communication, Stakeholder Management, Decision Making, Team Collaboration)
            - **Other Skills**: If a skill is found in the resume but **does not fit into any of the above categories**, include it under `"Other Skills"`.
        7. **Educational Qualifications**: Extract the highest level of education and any relevant degrees or diplomas, including the institution name.
            - If a category/bucket has no relevant skills, **exclude it from the output**.
            - If `"Other Skills"` is empty, **exclude it**.
            - If any field is not found, leave it as an empty string.

    ADDITIONAL FIELDS REQUIRED FOR THE SCREENING TABLE (MANDATORY OUTPUT KEYS):
        8. **Contact Number**: Extract phone number if present else empty string.
        9. **Email**: Extract email if present else empty string.
        10. **LinkedIn**: Extract LinkedIn profile URL if present else "Not available.
        11. **Visa Status**: If explicitly mentioned (e.g., "US Citizen", "Green Card", "H1B", etc.), output that text. If NOT mentioned, default to "Not available".
        12. **Relocation**: Output "Yes" only if the resume explicitly says willing/open to relocate (or equivalent). If NOT mentioned, default to "Not available".
        13. **Availability**: Any explicit availability info (e.g., timing, notice period). If NOT present, default to "Not available".
        14. **Video / In-person interview availability**: If explicitly mentioned, output "Yes" or the relevant wording. If NOT mentioned, default to "Not available".
        15. **Currently on Project**: Output "Yes" only if explicitly stated. If NOT mentioned, default to "Not available".
        16. **Any Certification**: Output "Yes" if certifications are explicitly listed, otherwise "Not available".
        17. **Education Entries (structured)**: Provide up to 3 entries with degree, institution and, year_range (if available). If missing, empty list of 1 row only.
        18. **Skill Matrix (optional)**: Only include skills with experience-in-years if explicitly stated. If not stated, keep this as an empty object.

    CRITICAL DEFAULT RULE (STRICT):
    For these fields, if the resume does NOT explicitly contain the information, set the value exactly to "Not available":
    - visa_status
    - relocation
    - availability
    - interview_availability
    - currently_on_project
    - any_certification

    DO NOT guess or infer. Only use what is explicitly present in the resume.

    Return the response in the following structured JSON format (ONLY JSON, no markdown/code fences):

    {{
        "name": "John Doe",
        "role": "Senior Software Engineer",
        "location": "San Francisco, CA",
        "experience": "8 years",
        "contact_number": "Not available",
        "email": "Not available",
        "linkedin": "Not available",
        "visa_status": "Not available",
        "relocation": "Not available",
        "availability": "Not available",
        "interview_availability": "Not available",
        "currently_on_project": "Not available",
        "any_certification": "Not available",
        "certifications": ["AWS Certified Solutions Architect", "PMP Certification"],
        "skills": {{
            "Programming Languages": ["Python", "Java"],
            "Databases": ["PostgreSQL", "MongoDB"],
            "Cloud Platforms": ["AWS", "Azure"],
            "Project Management Tools": ["Jira", "Trello"],
            "Reporting & Business Intelligence Tools": ["Tableau", "Power BI"],
            "Technologies & Frameworks": ["Salesforce", "React", "Django"],
            "DevOps & CI/CD Tools": ["Docker", "Kubernetes"],
            "Productivity Tools": ["Microsoft Word", "Microsoft PowerPoint"],
            "Soft Skills": ["Stakeholder Management", "Decision Making"],
            "Other Skills": ["Scrum Methodologies", "IT Governance"]
        }},
        "education": "Master of Science in Computer Science from Stanford University",
        "education_entries": [
            {{
                "degree": "",
                "institution": "",
                "year_range": ""
            }}
        ],
        "skill_matrix": {{}}
    }}

    Resume Text:
    ---
    {text}
    """

    response = _chat_completion_with_retry(
        messages=[
            {"role": "system", "content": "You are an HR expert analysing resumes."},
            {"role": "user", "content": prompt}
        ],
        temperature=0,
        run_dir=run_dir,
        call_name="parse_resume() — base resume JSON extraction"
    )

    llm_result = response.choices[0].message.content.strip()
    clean_result = re.sub(r"```(?:json)?", "", llm_result).strip()
    clean_result = clean_result.strip().lstrip("\ufeff").strip()

    try:
        parsed_data = json.loads(clean_result)
    except json.JSONDecodeError:
        print("Error parsing JSON. Response was:", clean_result)
        parsed_data = {}


    # Flag certification presence from raw resume text (even if list extraction fails)
    parsed_data["_cert_present"] = _cert_present_from_text(text)

    # Normalize any_certification to availability-only (no names)
    parsed_data["any_certification"] = _normalize_availability(_cert_availability(parsed_data))

    # If education_entries look incomplete but resume has multiple degrees in one line, enrich heuristically
    try:
        ee = parsed_data.get("education_entries") or []
        if not isinstance(ee, list):
            ee = []
        if len(ee) <= 1:
            heur = _heuristic_extract_education_entries(text)
            if len(heur) >= 2:
                parsed_data["education_entries"] = heur[:3]
    except Exception:
        pass

    return parsed_data


@traceable(name='extract_full_resume_via_llm')
def extract_full_resume_via_llm(resume_text: str, run_dir: Path) -> dict:
    """
    Stage A: Extract EVERYTHING of importance from the resume into a canonical structure
    with provenance, without mapping to template variables.
    Saves us from missing info when the template/prompt doesn't mention it.
    """
    chunks = _chunk_text_for_llm(resume_text, chunk_chars=14000, overlap=1200)
    all_parts = []

    for idx, chunk in enumerate(chunks, start=1):
        prompt = f"""
        You are extracting a resume into a COMPLETE structured JSON.
        Return ONLY valid JSON. No markdown, no commentary.

        ABSOLUTE RULES (CRITICAL):
        - Extract EVERYTHING of importance from the resume. Do not skip sections just because they are unfamiliar.
        - Do NOT merge content across sections in this stage.
        - Preserve provenance: every item must carry source_section_title + source_section_type + section_order.
        - Do NOT invent anything. If unsure, include the text as an item with kind="other" rather than skipping.
        - Keep the order of sections as in the resume (top to bottom).
        - For each section, capture both paragraphs and bullets as separate items.
        - For experience blocks: extract company/client, role/title, dates/duration, and bullets under that role (as items).

        SECTION TYPING GUIDANCE:
        - summary: QUALIFICATIONS / SUMMARY / PROFESSIONAL SUMMARY / CAREER HIGHLIGHTS / PROFILE / TECHNICAL SUMMARY (if it is written like a summary)
        - skills: PROFESSIONAL SKILLS / TECHNICAL EXPERTISE / SKILLS / CORE COMPETENCIES / TOOLS
        - experience/projects: PROFESSIONAL EXPERIENCE / WORK EXPERIENCE / PROJECT EXPERIENCE
        - education: EDUCATION
        - certifications: CERTIFICATIONS (including cert lists under EDUCATION/CERTIFICATIONS)

        OUTPUT MUST MATCH THIS SCHEMA (same keys; you may expand arrays, but do not add new top-level keys):
        {json.dumps(FULL_SCHEMA, ensure_ascii=False)}

        RESUME_TEXT:
        ---
        {chunk}
        ---
        """

        response = _chat_completion_with_retry(
            messages=[
                {"role": "system", "content": "You output strictly valid JSON only."},
                {"role": "user", "content": prompt}
            ],
            temperature=0,
            run_dir=run_dir,
            call_name=f"extract_full_resume_via_llm() — chunk {idx} of {len(chunks)}"
        )

        content = response.choices[0].message.content.strip()
        content = re.sub(r"```(?:json)?", "", content).strip()
        content = content.strip().lstrip("\ufeff").strip()

        try:
            part = json.loads(content)
        except json.JSONDecodeError:
            part = {}

        # Defensive normalization
        if isinstance(part, dict):
            all_parts.append(part)
    
    full = _merge_extracted_full(all_parts)

    # Defensive normalization (keep your existing code)
    if "sections" not in full or not isinstance(full.get("sections"), list):
        full["sections"] = []
    return full


@traceable(name='to_firm_json')
def to_firm_json(base_json: dict, resume_text: str, run_dir: Path, extracted_full: dict | None = None) -> dict:
    """Convert parsed resume JSON + resume text into Company version firm-resume JSON for DOCX placeholders."""
    NOT_AVAILABLE_LOCAL = "Not available"
    resume_text = _truncate_for_llm(resume_text)

    # NOTE: This schema MUST match the placeholders used in Company_Template.docx
    firm_schema = {
        "name": "",
        "candidate_name": "",   # kept for backward compatibility (some templates used this)
        "email": NOT_AVAILABLE_LOCAL,
        "contact_number": NOT_AVAILABLE_LOCAL,
        "linkedin": "",
        "title": "",
        "location": "",
        "education_entries": [
            {"degree": "", "institution": "", "year_range": ""}
        ],
        "certifications": [],  # list of strings (can be empty)
        "experience_summary_lines": [""],
        "technical_skills": [
            {"category": "", "items": ""}  # items MUST be a single comma-separated STRING
        ],
        "projects": [
            {
                "client": "",
                "summary": "",
                "tools_environment": "",
                "role": "",
                "duration": ""
            }
        ]
    }

    prompt = f"""
    You are mapping a FULL extracted resume JSON into Company firm-resume JSON (template variables).
    Return ONLY valid JSON. No markdown, no commentary, no extra keys.

    IMPORTANT: Stage-A extraction already captured everything with provenance.
    Your job now is ONLY to map and merge *similar* sections into the template, while enforcing strict routing rules.

    INPUTS:
    1) BASE_JSON: basic parsed fields (name/role/location/skills/certs/education etc.)
    2) EXTRACTED_FULL: complete resume extraction with sections/items/provenance
    3) RESUME_TEXT: raw text for verification only (do not use it to invent missing things)

    NON-NEGOTIABLE RULES:
    - Do NOT invent companies, roles, dates, tools, or bullets.
    - Do NOT move content across unrelated areas.
    - "Experience Summary" (experience_summary_lines) must come ONLY from summary-like sections:
    titles/types like QUALIFICATIONS, SUMMARY, PROFESSIONAL SUMMARY, PROFILE, CAREER HIGHLIGHTS, TECHNICAL SUMMARY (ONLY if it reads like a summary).
    It must NOT include any bullets that belong to a role/company under PROFESSIONAL EXPERIENCE / WORK EXPERIENCE.
    - "Client/Role summaries" (projects[].summary) must come ONLY from bullets/text that are truly under that role/company block in the extracted data.
    If the resume does not present role bullets, do not fabricate them and do not borrow from other sections.
    - Template flexibility requirement:
    If the resume has multiple summary-like sections (e.g., QUALIFICATIONS + CAREER HIGHLIGHTS + TECHNICAL SUMMARY),
    you may MERGE them into experience_summary_lines, but ONLY because they are summary-like.
    Similarly, if the resume has multiple skill-like sections (e.g., PROFESSIONAL SKILLS + TECHNICAL EXPERTISE),
    you may MERGE them into technical_skills buckets.

    COVERAGE GUARANTEE (CRITICAL):
    - Do not drop important content.
    - Every important summary-like section item should appear in either:
    a) experience_summary_lines, OR
    b) technical_skills (if it is clearly skills/tools list), OR
    c) projects/tools_environment/projects.summary (if it is clearly role experience content), OR
    d) certifications/education_entries
    - If something is important but cannot fit the template cleanly, include it in:
    experience_summary_lines as additional lines ONLY IF it is summary-like;
    otherwise include it inside the most relevant project summary where it truly belongs.
    Never violate the routing rules above.

    OUTPUT JSON MUST MATCH EXACTLY THIS SCHEMA (no extra keys):
    {json.dumps(firm_schema, ensure_ascii=False)}

    BASE_JSON:
    {json.dumps(base_json, ensure_ascii=False)}

    EXTRACTED_FULL:
    {json.dumps(extracted_full or {}, ensure_ascii=False)}

    RESUME_TEXT (verification only):
    ---
    {resume_text}
    ---
    """

    response = _chat_completion_with_retry(
        messages=[
            {"role": "system", "content": "You output strictly valid JSON only."},
            {"role": "user", "content": prompt}
        ],
        temperature=0,
        run_dir=run_dir,
        call_name="to_firm_json() — map extracted data to Company template JSON"
    )

    content = response.choices[0].message.content.strip()
    firm_json = json.loads(content)

    # Fail-safe: if we somehow got too few project bullets, retry mapping once with higher chunk safety
    try:
        proj = firm_json.get("projects", []) or []
        total_lines = sum(len(p.get("summary", []) or []) for p in proj if isinstance(p, dict))
        if total_lines < 8 and len(proj) >= 2:  # tune threshold as you like
            firm_json = to_firm_json(parsed_data, resume_text, run_dir, extracted_full=extracted_full)
            firm_json = _sanitize_for_docx(firm_json)
    except Exception:
        pass

    # ---- sanitize technical_skills so s.items is always a STRING ----
    def _to_items_string(x):
        if x is None:
            return ""
        if isinstance(x, str):
            return x.strip()
        if isinstance(x, list):
            return ", ".join(_to_items_string(i) for i in x if _to_items_string(i))
        if isinstance(x, dict):
            parts = []
            for _, v in x.items():
                sv = _to_items_string(v)
                if sv:
                    parts.append(sv)
            return ", ".join(parts)
        return str(x).strip()

    skills = firm_json.get("technical_skills", []) or []
    fixed_skills = []
    for s in skills:
        if not isinstance(s, dict):
            continue
        category = str(s.get("category", "") or "").strip()
        items_str = _to_items_string(s.get("items", ""))
        fixed_skills.append({"category": category, "items": items_str})
    firm_json["technical_skills"] = fixed_skills
    # ---------------------------------------------------------------


    # ---- GUARANTEE header fields for template ----
    # name
    name = str(firm_json.get("name", "") or "").strip()
    if not name:
        name = str(firm_json.get("candidate_name", "") or "").strip()
    if not name:
        name = str(base_json.get("name", "") or "").strip()
    firm_json["name"] = name
    firm_json["candidate_name"] = name  # keep both aligned

    # title/location
    if not str(firm_json.get("title", "") or "").strip():
        firm_json["title"] = str(base_json.get("role", "") or "").strip()
    if not str(firm_json.get("location", "") or "").strip():
        firm_json["location"] = str(base_json.get("location", "") or "").strip()

    # email/contact_number/linkedin (for "Name | Phone | Email" line in template)
    email = str(firm_json.get("email", "") or "").strip()
    if not email:
        email = str(base_json.get("email", "") or "").strip()
    firm_json["email"] = email if email else NOT_AVAILABLE_LOCAL

    phone = str(firm_json.get("contact_number", "") or "").strip()
    if not phone:
        phone = str(base_json.get("contact_number", "") or "").strip()
    firm_json["contact_number"] = phone if phone else NOT_AVAILABLE_LOCAL

    linkedin = str(firm_json.get("linkedin", "") or "").strip()
    if not linkedin:
        linkedin = str(base_json.get("linkedin", "") or "").strip()
    firm_json["linkedin"] = linkedin if linkedin else NOT_AVAILABLE_LOCAL

    # education_entries (prefer structured from base_json if LLM omitted)
    ee = firm_json.get("education_entries")
    if not isinstance(ee, list) or not ee:
        ee = base_json.get("education_entries")
    if not isinstance(ee, list) or not ee:
        ee = [{"degree": "", "institution": "", "year_range": ""}]
    firm_json["education_entries"] = ee[:3]

    # ---- Clean education_entries so template doesn't need if() ----
    clean_edu = []
    for e in firm_json.get("education_entries", []) or []:
        if not isinstance(e, dict):
            continue
        deg = (e.get("degree") or "").strip()
        inst = (e.get("institution") or "").strip()
        yr = (e.get("year_range") or "").strip()
        if not (deg or inst or yr):
            continue
        clean_edu.append({"degree": deg, "institution": inst, "year_range": yr})

    firm_json["education_entries"] = clean_edu
    # -------------------------------------------------------------


    # certifications (prefer extracted list from base_json if LLM omitted)
    certs = firm_json.get("certifications")
    if not isinstance(certs, list):
        certs = base_json.get("certifications")
    if not isinstance(certs, list):
        certs = []
    firm_json["certifications"] = [str(c).strip() for c in certs if str(c).strip()]
    if not firm_json["certifications"]:
        firm_json["certifications"] = ["Not available"]

    def _ensure_lines(x, *, max_lines: int = 40, max_chars_per_line: int = 170):
        """
        Convert summary into list of bullet-sized lines without overfitting to verbs.
        Strategy:
        1) split on explicit separators (newlines/bullets)
        2) split on sentence-ish boundaries (. ! ? ; :)
        3) if still huge, chunk by length at word boundaries
        """
        if x is None:
            return []
        if isinstance(x, list):
            out = []
            for item in x:
                out.extend(_ensure_lines(item, max_lines=max_lines, max_chars_per_line=max_chars_per_line))
            # de-dupe empties
            return [s for s in out if s]

        s = str(x).strip()
        if not s:
            return []

        # Normalize whitespace
        s = str(x).strip()
        # Keep newlines; just normalize spaces/tabs (not \n)
        s = re.sub(r"[ \t]+", " ", s)
        s = re.sub(r"\n{3,}", "\n\n", s).strip()

        # 1) explicit separators first
        # Split on:
        # - real bullet symbols
        # - OR line breaks that look like a new bullet/entry (indent + capital/number)
        parts = re.split(
            r"[•▪]|(?:\r?\n)+(?=\s*(?:[-–—*]|\d+\)|\d+\.|[A-Z]))",
            s
        )
        parts = [p.strip(" -–—•▪\t").strip() for p in parts if p and p.strip()]

        # --- NEW: merge continuation lines caused by hard line breaks in source DOCX/PDF ---
        merged = []
        for p in parts:
            if not merged:
                merged.append(p)
                continue

            prev = merged[-1].rstrip()

            # If next chunk starts lowercase OR prev doesn't end a sentence, treat as same bullet continuation
            if (p[:1].islower()) or (prev and prev[-1] not in ".!?"):
                merged[-1] = (prev + " " + p).strip()
            else:
                merged.append(p)

        parts = merged
        # -------------------------------------------------------------------------------

        # 2) sentence-ish split if we still have one big chunk
        if len(parts) <= 1:
            # split after punctuation when followed by space and next token looks like a new sentence-ish start
            parts = re.split(r"(?<=[.!?])\s+|;\s+|:\s+", s)
            parts = [p.strip(" -–—•▪\t").strip() for p in parts if p and p.strip()]

        # 3) IMPORTANT: do NOT split long bullets by length.
        # Word will wrap within the same bullet paragraph automatically.
        final = [p for p in parts if p.strip()]
        return final[:max_lines]
    
    def _pair_compress_bullets(lines: list[str], *, max_len_each: int = 90) -> list[str]:
        """
        Conservative space-saver:
        - Merge consecutive short bullets into one line: "a, b"
        - Leaves long/complex bullets untouched.
        """
        if not isinstance(lines, list):
            return lines

        cleaned = [str(x).strip() for x in lines if str(x or "").strip()]
        out = []
        i = 0

        while i < len(cleaned):
            a = cleaned[i]
            b = cleaned[i + 1] if i + 1 < len(cleaned) else None

            def _ok(x: str) -> bool:
                return len(x) <= max_len_each

            def _ends_strong(x: str) -> bool:
                return x.endswith((".", "!", "?", ":"))

            if b and _ok(a) and _ok(b) and (not _ends_strong(a)) and (not _ends_strong(b)):
                a2 = a.rstrip(" ,;")
                b2 = b.lstrip(" ,;")
                out.append(f"{a2}, {b2}")
                i += 2
            else:
                out.append(a)
                i += 1

        return out

    # ---- Normalize projects: ensure client + tools_environment exist and have fallbacks ----
    projects = firm_json.get("projects", []) or []
    fixed_projects = []
    for p in projects:
        if not isinstance(p, dict):
            continue
        client_name = str(p.get("client", "") or "").strip()
        if not client_name:
            client_name = NOT_AVAILABLE_LOCAL
        p["client"] = client_name

        tools_env = str(p.get("tools_environment", "") or p.get("tools", "") or p.get("tools_and_environment", "") or "").strip()
        p["tools_environment"] = tools_env  # can be empty string

        # Keep required keys present
        p["summary"] = _pair_compress_bullets(
            _ensure_lines(p.get("summary"), max_lines=60),
            max_len_each=90
        )
        p["role"] = p.get("role", "") or ""
        p["duration"] = p.get("duration", "") or ""
        fixed_projects.append(p)

    firm_json["projects"] = fixed_projects or firm_schema["projects"]

    return firm_json


@traceable(name='apply_excel_styling')
def apply_excel_styling(file_path: str) -> None:
    workbook = load_workbook(file_path)
    sheet = workbook.active

    for col in sheet.columns:
        max_length = 0
        col_letter = col[0].column_letter
        col_name = sheet.cell(row=1, column=col[0].column).value

        for cell in col:
            if cell.value:
                max_length = max(max_length, len(str(cell.value)))
                if col_name == "Technical Skills":
                    cell.alignment = Alignment(wrap_text=True)
                    sheet.column_dimensions[col_letter].width = 100
                elif col_name == "Technical Certifications":
                    cell.alignment = Alignment(wrap_text=True)
                    sheet.column_dimensions[col_letter].width = 30
                else:
                    sheet.column_dimensions[col_letter].width = max_length + 1

    for cell in sheet[1]:
        cell.font = Font(bold=True)
        cell.alignment = Alignment(horizontal="center")

    workbook.save(file_path)


@traceable(name='profile_summary_excel')
def profile_summary_excel(out_xlsx_path: str, parsed_data: dict) -> str:
    """
    Creates a single-resume Excel output in the given path.
    """
    skills_dict = parsed_data.get("skills", {}) or {}
    formatted_skills = [f"{category}: {', '.join(skills)}"
                        for category, skills in skills_dict.items()
                        if isinstance(skills, list) and skills]
    skills_text = "\n".join(formatted_skills) if formatted_skills else ""

    certifications_text = _cert_availability(parsed_data)

    df = pd.DataFrame([{
        "Name": parsed_data.get("name", ""),
        "Role": parsed_data.get("role", ""),
        "Location": parsed_data.get("location", ""),
        "Years of Experience": parsed_data.get("experience", ""),
        "Technical Certifications": certifications_text,
        "Technical Skills": skills_text,
        "Educational Qualifications": parsed_data.get("education", "")
    }])

    with pd.ExcelWriter(out_xlsx_path, engine="openpyxl") as writer:
        df.to_excel(writer, index=False)

    apply_excel_styling(out_xlsx_path)
    return out_xlsx_path


# Resume Fields Table in UI
UI_TABLE_TEMPLATE_XLSX = "templates/ui table.xlsx"  # keep alongside app, or change path if needed

def _compact_two_col_template(df: pd.DataFrame) -> pd.DataFrame:
    """
    1) Remove accidental header-as-row ("Particular","Details")
    2) Replace NaN with ""
    3) Compact rows where Particular is blank but Details has text:
       append Details to previous row (newline) and drop that row.
    """
    df = df.copy()

    # Ensure exactly 2 columns
    df = df.iloc[:, :2].copy()
    df.columns = ["Particular", "Details"]

    # Drop row if it is the header repeated as data
    if len(df) > 0:
        first_p = str(df.iloc[0]["Particular"]).strip().lower()
        first_d = str(df.iloc[0]["Details"]).strip().lower()
        if first_p == "particular" and first_d == "details":
            df = df.iloc[1:].reset_index(drop=True)

    # Replace NaN/None with ""
    df["Particular"] = df["Particular"].fillna("").astype(str)
    df["Details"] = df["Details"].fillna("").astype(str)

    # Compact: if Particular empty but Details not empty => append to previous Details
    rows_to_drop = []

    for i in range(len(df)):
        p = df.at[i, "Particular"].strip()
        d = df.at[i, "Details"].strip()
        if i > 0 and p == "" and d != "":
            prev_particular = str(df.at[i - 1, "Particular"]).strip().lower()

            # ❌ DO NOT compact education rows
            if "education" in prev_particular or "educational qualification" in prev_particular:
                continue

            prev = df.at[i - 1, "Details"].rstrip()
            df.at[i - 1, "Details"] = (prev + "\n" + d).strip() if prev else d
            rows_to_drop.append(i)

    if rows_to_drop:
        df = df.drop(index=rows_to_drop).reset_index(drop=True)

    return df


def load_ui_table_template(xlsx_path: str, sheet_name=0) -> pd.DataFrame:
    """
    Loads a 2-column template sheet (Particular, Details) and normalizes it.
    sheet_name can be index or sheet name.
    """
    raw = pd.read_excel(xlsx_path, sheet_name=sheet_name, header=None)
    return _compact_two_col_template(raw)


NOT_AVAILABLE = "Not available"

# Detect certifications even when LLM doesn't list them (e.g., "AWS certified")
_CERT_PRESENCE_RE = re.compile(r"(?i)\b(certification|certifications|certified)\b")

def _cert_present_from_text(resume_text: str) -> bool:
    if not resume_text:
        return False
    return bool(_CERT_PRESENCE_RE.search(str(resume_text)))

def _cert_availability(parsed_data: dict) -> str:
    """
    Availability-only certifications:
    - "Available" if ANY certification is present (either extracted list OR detected by text flag)
    - else "Not available"
    """
    certs = parsed_data.get("certifications")
    has_list = isinstance(certs, list) and any(str(c).strip() for c in certs or [])
    has_flag = bool(parsed_data.get("_cert_present", False))
    return "Available" if (has_list or has_flag) else NOT_AVAILABLE


def _required_default(v) -> str:
    """If missing/empty -> 'Not available', else keep string value."""
    if v is None:
        return NOT_AVAILABLE
    s = str(v).strip()
    return s if s else NOT_AVAILABLE


def build_ui_table_from_parsed(template_df: pd.DataFrame, parsed_data: dict) -> pd.DataFrame:
    """
    Returns a dataframe that matches the template, but fills the 'Details' column
    from parsed_data (parsed_resume.json).
    """
    df = template_df.copy()
    df["Details"] = ""  # IMPORTANT: clear template example values

    # --- pull from parsed_data (existing keys + new keys we add to prompt) ---
    # Education: allow structured list; fallback to existing 'education' string
    education_fallback = parsed_data.get("education", "") or ""

    # Skill Matrix: allow structured dict; fallback minimal values
    skill_matrix = parsed_data.get("skill_matrix", {}) or {}

    # --- mapping for template Particular values -> Details ---
    def _val(x):
        s = "" if x is None else str(x).strip()
        return s if s else NOT_AVAILABLE

    PARTICULAR_TO_JSON_KEY = {
        "full name of the candidate": "name",
        "location": "location",
        "contact number": "contact_number",
        "email id": "email",
        "linkedin": "linkedin",
        "visa status (along with validity details)": "visa_status",
        "relocation": "relocation",
        "availability": "availability",
        "video / in-person interview availability (yes/no)": "interview_availability",
        "currently on project": "currently_on_project",
        "any certification": "any_certification",
        "technical certifications": "any_certification",
        }

    def _norm(s: str) -> str:
        return re.sub(r"\s+", " ", str(s or "")).strip().lower()

    for idx, row in df.iterrows():
        p = _norm(row.get("Particular", ""))
        if p in PARTICULAR_TO_JSON_KEY:
            key = PARTICULAR_TO_JSON_KEY[p]
            df.at[idx, "Details"] = _val(parsed_data.get(key))
    
    # Education block filling (robust + supports multiple entries even if template has only one row)
    # Education: show each degree on its own row (no single-line clumsy output)
    edu_mask = df["Particular"].astype(str).str.contains("Educational qualification", case=False, na=False)
    edu_indices = df.index[edu_mask].tolist()

    if edu_indices:
        start_i = edu_indices[0]

        # Build education display lines from structured entries
        lines = []
        for e in (parsed_data.get("education_entries") or [])[:5]:
            if not isinstance(e, dict):
                continue
            deg = (e.get("degree") or "").strip()
            inst = (e.get("institution") or "").strip()
            yr = (e.get("year_range") or "").strip()
            if not deg:
                continue
            parts = [deg]
            if inst:
                parts.append(inst)
            if yr:
                parts.append(yr)
            lines.append(" - ".join(parts))
        
        # Post-clean: split clumsy combined education lines
        cleaned_lines = []
        for ln in lines:
            chunks = re.split(
                r"(?i)(?<=\b(19\d{2}|20\d{2})\b)\s+(?=(bachelor|master|phd|doctorate|b\.|m\.))",
                ln
            )

            buf = ""
            for c in chunks:
                if not c:
                    continue
                if _DEGREE_MARKERS_RE.search(c) and buf.strip():
                    cleaned_lines.append(buf.strip(" -"))
                    buf = c
                else:
                    buf = (buf + " " + c).strip() if buf else c

            if buf.strip():
                cleaned_lines.append(buf.strip(" -"))

        lines = [x for x in cleaned_lines if x.strip()]

        if not lines:
            edu_str = (parsed_data.get("education") or "").strip()
            lines = [edu_str] if edu_str else [NOT_AVAILABLE]

        # Identify "slots" under the education header: consecutive rows where Particular is blank
        # Treat the education header row itself as the first slot
        slot_idxs = [start_i]
        j = start_i + 1
        while j < len(df) and str(df.at[j, "Particular"] or "").strip() == "":
            slot_idxs.append(j)
            j += 1

        if slot_idxs:
            # Fill existing slots
            for k, idx_row in enumerate(slot_idxs):
                if k < len(lines):
                    df.at[idx_row, "Details"] = lines[k]
                else:
                    df.at[idx_row, "Details"] = ""
            # If more lines than slots, insert additional blank rows right after the last slot
            if len(lines) > len(slot_idxs):
                insert_at = slot_idxs[-1] + 1
                extra = pd.DataFrame(
                    [{"Particular": "", "Details": ln} for ln in lines[len(slot_idxs):]],
                    columns=["Particular", "Details"]
                )
                df = pd.concat([df.iloc[:insert_at], extra, df.iloc[insert_at:]], ignore_index=True)
        else:
            # If template has no slots, still avoid clumsy header-fill.
            # Put the first education line in the header's Details ONLY if it's Not available.
            if lines and lines[0] == NOT_AVAILABLE:
                df.at[start_i, "Details"] = NOT_AVAILABLE
            else:
                # Insert rows below header for each education line
                insert_at = start_i + 1
                extra = pd.DataFrame(
                    [{"Particular": "", "Details": ln} for ln in lines],
                    columns=["Particular", "Details"]
                )
                df = pd.concat([df.iloc[:insert_at], extra, df.iloc[insert_at:]], ignore_index=True)

    # Skill matrix dynamic keys that appear in template (like Azure AI studio etc.)
    # If prompt returns skill_matrix with those exact keys, we fill them.
    # Fill any skill_matrix keys where the template "Particular" matches exactly
    # Certifications must be availability-only (never show names)
    cert_display = "Available" if _cert_availability(parsed_data) != NOT_AVAILABLE else NOT_AVAILABLE

    for i in range(len(df)):
        p_norm = _norm(df.at[i, "Particular"])
        if re.search(r"\bcert", p_norm):   # catches certification/certified/certifications/certificate
            df.at[i, "Details"] = cert_display

    return df


# ------------------------- UI -------------------------

logo_path = os.path.join("Company-logo.png")
st.image(logo_path, width=250)

st.markdown("<H1 style='color: #e2a244; text-align: left;'>Resume Standardizer | Company</H1>", unsafe_allow_html=True)
st.markdown("<h3 style='color: #0e3b79; text-align: left, font-size=30px;'>AI-Driven Resume Standardization for Company</h3>", unsafe_allow_html=True)
st.markdown(
    "<p style='color: #0e3b79; text-align: left;'>"
    "Upload PDF/ Word Document resumes to extract structured candidate information and generate resumes aligned with Company’s official formatting standards."
    "<br><b>Note:</b> You can upload up to <b>3</b> resumes at a time."
    "</p>",
    unsafe_allow_html=True
)
st.markdown("---")


# --- Upload controls: compact uploader + matched-height clear button ---
col_up, col_clear = st.columns([5, 1.4], vertical_alignment="center")

with col_up:
    uploaded_files = st.file_uploader(
        "Upload Resume PDFs or Word Documents",
        type=["pdf", "docx"],
        accept_multiple_files=True,
        key=st.session_state.uploader_key
    )

with col_clear:
    st.markdown('<div class="clear-tile">', unsafe_allow_html=True)
    if st.button("Clear\nResumes", use_container_width=True, key="clear_resumes_btn"):
        st.session_state.uploader_key = str(uuid.uuid4())
        st.session_state.run_map = {}
        st.session_state.processed = set()
        st.session_state.batch_llm_usage = []
        st.rerun()
    st.markdown("</div>", unsafe_allow_html=True)


if uploaded_files:
    st.toast(f"Processing {len(uploaded_files)} resumes, please wait...")
    # Reset batch totals so LLM USAGE is calculated only for THIS upload run
    _reset_batch_llm_usage()
    progress_bar = st.progress(0)
    total_files = len(uploaded_files)

    ensure_dir(OUTPUT_ROOT)

    for i, uploaded_file in enumerate(uploaded_files):
        # --- Per-resume status box (updates in-place: Processing -> Processed) ---
        status_box = st.empty()
        status_box.info(f"⏳ Processing: {uploaded_file.name}")

        # Stable id for this upload (prevents duplicate runs when Streamlit reruns)
        file_bytes = uploaded_file.getvalue()
        file_id = hashlib.md5(file_bytes).hexdigest()[:12]

        if file_id not in st.session_state.run_map:
            st.session_state.run_map[file_id] = create_run_dir(uploaded_file.name)
        run_dir = st.session_state.run_map[file_id]

        # If already processed in this session, reuse outputs but still render UI + download
        if file_id in st.session_state.processed:
            status_box.success(f"✅ Processed: {uploaded_file.name}")

            # Load cached parsed JSON if available
            parsed_json_path = run_dir / "parsed_resume.json"
            parsed_data = {}
            if parsed_json_path.exists():
                try:
                    parsed_data = json.loads(parsed_json_path.read_text(encoding="utf-8"))
                except Exception:
                    parsed_data = {}

            # Render Candidate Details table again
            try:
                candidate_df = load_ui_table_template(UI_TABLE_TEMPLATE_XLSX, sheet_name=0)
                filled_candidate_df = build_ui_table_from_parsed(candidate_df, parsed_data)
                st.markdown("### Candidate Details")
                st.dataframe(filled_candidate_df, use_container_width=True, hide_index=True)
                render_copy_button_for_df(filled_candidate_df, button_label="📋 Copy Candidate Details table", unique_key=f"cand_{file_id}")
            except Exception as e:
                st.warning(f"Could not render candidate tables for {uploaded_file.name}: {e}")

            # Render Skill Matrix again if firm JSON exists
            try:
                firm_json_path = run_dir / "firm_resume.json"
                if firm_json_path.exists():
                    firm_json = json.loads(firm_json_path.read_text(encoding="utf-8"))
                    skill_matrix_df = build_dynamic_skill_matrix(firm_json)
                    st.markdown("### Skill Matrix")
                    if skill_matrix_df.empty:
                        st.info("Skill Matrix: Not available")
                    else:
                        st.dataframe(skill_matrix_df, use_container_width=True, hide_index=True)
                        render_copy_button_for_df(skill_matrix_df, button_label="📋 Copy Skill Matrix table", unique_key=f"skill_{file_id}")
            except Exception as e:
                st.warning(f"Skill Matrix generation failed for {uploaded_file.name}: {e}")

            # Download button again if DOCX exists
            safe_stem = safe_name(uploaded_file.name)
            firm_docx_path = run_dir / f"Company_{safe_stem}.docx"
            non_docx_path = run_dir / f"Non_Company_{safe_stem}.docx"
            if firm_json_path.exists():
                try:
                    firm_json = json.loads(firm_json_path.read_text(encoding="utf-8"))
                    render_firm_resume(
                        template_path=NON_TEMPLATE_DOCX,
                        firm_json=firm_json,
                        out_docx_path=str(non_docx_path)
                    )
                    postprocess_docx_spacing(str(non_docx_path))
                except Exception:
                    pass
            if firm_docx_path.exists():
                with open(firm_docx_path, "rb") as f:
                    st.download_button(
                        label=f"Download Company Version (DOCX) - {firm_docx_path.name}",
                        data=f,
                        file_name=firm_docx_path.name,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
            if non_docx_path and non_docx_path.exists():
                with open(non_docx_path, "rb") as f:
                    st.download_button(
                        label=f"Download Non Company Version (DOCX) - {non_docx_path.name}",
                        data=f,
                        file_name=non_docx_path.name,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
            continue
        # Extract text directly from uploaded PDF bytes (do NOT save PDF to output folder)
        file_ext = Path(uploaded_file.name).suffix.lower()

        if file_ext == ".pdf":
            resume_text = load_pdf_bytes(file_bytes)
        elif file_ext == ".docx":
            resume_text = load_docx_bytes(file_bytes)
        else:
            st.warning(f"Unsupported file type: {uploaded_file.name}")
            continue
        
        # Parse + save JSON to outputs folder (no longer shown in UI)
        parsed_data = parse_resume(resume_text, run_dir)

        # --- NEW: Stage-A full extraction (save to output folder) ---
        extracted_full = extract_full_resume_via_llm(resume_text, run_dir)
        extracted_full_path = run_dir / "extracted_full.json"
        extracted_full_path.write_text(
            json.dumps(extracted_full, ensure_ascii=False, indent=2),
            encoding="utf-8"
        )
        # -----------------------------------------------------------

        # ---- DB INSERT: store uploaded resume filename (one-time per file_id) ----
        try:
            insert_resume_upload(uploaded_file.name)  # resume_name = uploaded doc name
        except Exception as e:
            st.warning(f"DB insert failed for {uploaded_file.name}: {e}")
        # ------------------------------------------------------------------------

        parsed_json_path = run_dir / "parsed_resume.json"
        parsed_json_path.write_text(
            json.dumps(parsed_data, ensure_ascii=False, indent=2),
            encoding="utf-8"
        )

        # --------------- LLM Call #2: Enrich UI-focused fields (education, screening fields) ---------------
        try:
            candidate_template_df = load_ui_table_template(UI_TABLE_TEMPLATE_XLSX, sheet_name=0)
            enrich = enrich_ui_fields_via_llm(parsed_data, resume_text, run_dir, template_df=candidate_template_df)
            if isinstance(enrich, dict) and enrich:
                parsed_data.update(enrich)

                # Enforce availability-only certifications (no names)
                parsed_data["any_certification"] = _normalize_availability(_cert_availability(parsed_data))
                # Heuristic education fix if multiple degrees are on one line
                try:
                    ee2 = parsed_data.get("education_entries") or []
                    if not isinstance(ee2, list):
                        ee2 = []
                    if len(ee2) <= 1:
                        heur2 = _heuristic_extract_education_entries(resume_text)
                        if len(heur2) >= 2:
                            parsed_data["education_entries"] = heur2[:3]
                except Exception:
                    pass

                parsed_json_path.write_text(json.dumps(parsed_data, ensure_ascii=False, indent=2), encoding="utf-8")
        except Exception:
            pass
        # -----------------------------------------------------------------------------------------------

        # ---------------- UI Tables (Read-only) ----------------
        try:
            # 1) Candidate Details table from Sheet 0
            candidate_df = load_ui_table_template(UI_TABLE_TEMPLATE_XLSX, sheet_name=0)
            filled_candidate_df = build_ui_table_from_parsed(candidate_df, parsed_data)

            st.markdown("### Candidate Details")
            st.dataframe(filled_candidate_df, use_container_width=True, hide_index=True)
            render_copy_button_for_df(filled_candidate_df, button_label="📋 Copy Candidate Details table", unique_key=f"cand_{file_id}")

        except Exception as e:
            st.warning(f"Could not render candidate tables for {uploaded_file.name}: {e}")
        # ------------------------------------------------------

        # Firm version generation (DOCX only) inside the same run folder
        firm_docx_path = None
        firm_json_path = None

        safe_stem = safe_name(uploaded_file.name)
        firm_docx_path = run_dir / f"Company_{safe_stem}.docx"

        try:
            firm_json = to_firm_json(parsed_data, resume_text, run_dir, extracted_full=extracted_full)
            firm_json = _sanitize_for_docx(firm_json)

            # --------- Dynamic Skill Matrix (Read-only) ----------
            try:
                skill_matrix_df = build_dynamic_skill_matrix(firm_json)
                st.markdown("### Skill Matrix")
                if skill_matrix_df.empty:
                    st.info("Skill Matrix: Not available")
                else:
                    st.dataframe(skill_matrix_df, use_container_width=True, hide_index=True)
                    render_copy_button_for_df(skill_matrix_df, button_label="📋 Copy Skill Matrix table", unique_key=f"skill_{file_id}")
            except Exception as e:
                st.warning(f"Skill Matrix generation failed for {uploaded_file.name}: {e}")
            # -----------------------------------------------------

            firm_json_path = run_dir / "firm_resume.json"
            firm_json_path.write_text(
                json.dumps(firm_json, ensure_ascii=False, indent=2),
                encoding="utf-8"
            )

            render_firm_resume(
                template_path=TEMPLATE_DOCX,
                firm_json=firm_json,
                out_docx_path=str(firm_docx_path)
            )
            postprocess_docx_spacing(str(firm_docx_path))

            # --- Non-Company external version (DOCX) ---
            non_docx_path = run_dir / f"Non_Company_{safe_stem}.docx"

            render_firm_resume(
                template_path=NON_TEMPLATE_DOCX,
                firm_json=firm_json,
                out_docx_path=str(non_docx_path)
            )
            postprocess_docx_spacing(str(non_docx_path), keep_space_before_headers=True)

            if non_docx_path.stat().st_size < 5000:
                raise RuntimeError(
                    f"Generated NON DOCX too small ({non_docx_path.stat().st_size} bytes). "
                    "Template or data issue."
                )
            # ------------------------------------------

            # Basic validation: template/render failures often create tiny DOCX files
            if firm_docx_path.stat().st_size < 5000:
                raise RuntimeError(
                    f"Generated DOCX too small ({firm_docx_path.stat().st_size} bytes). "
                    "Template or data issue."
                )
            
        except Exception as e:
            firm_docx_path = None
            st.warning(f"Firm Version (DOCX) generation failed for {uploaded_file.name}: {e}")

        finally:
            # Never keeps any FirmVersion PDF in the run folder
            try:
                for p in list(run_dir.glob("FirmVersion_*.pdf")) + list(run_dir.glob("Company_*.pdf")):
                    try:
                        p.unlink()
                    except Exception:
                        pass
            except Exception:
                pass

        # Create per-resume Excel inside the same run folder
        excel_path = run_dir / "Resume_Output.xlsx"
        profile_summary_excel(str(excel_path), parsed_data)

        # Download button (DOCX only)
        if firm_docx_path and firm_docx_path.exists():
            with open(firm_docx_path, "rb") as f:
                st.download_button(
                    label=f"Download Company Version (DOCX) - {firm_docx_path.name}",
                    data=f,
                    file_name=firm_docx_path.name,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
        
        # Download Non-Company version (DOCX only)
        if non_docx_path and non_docx_path.exists():
            with open(non_docx_path, "rb") as f:
                st.download_button(
                    label=f"Download Non Company Version (DOCX) - {non_docx_path.name}",
                    data=f,
                    file_name=non_docx_path.name,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

        # Excel is generated and saved in the output folder, but NOT shown as a download button in the UI.
        st.session_state.processed.add(file_id)

        status_box.success(f"✅ Processed: {uploaded_file.name}")

        progress_bar.progress((i + 1) / total_files)

    progress_bar.progress(1.0)
    st.success("✅ All resumes processed.")
    _print_batch_llm_totals()